import os
import tempfile
from typing import List, Optional
from pathlib import Path

from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from docx import Document as DocxDocument


class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def load_pdf(self, file_path: str) -> List[Document]:
        loader = PyPDFLoader(file_path)
        return loader.load()
    
    def load_docx(self, file_path: str) -> List[Document]:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return [Document(page_content=text, metadata={"source": file_path})]
    
    def process_uploaded_file(self, uploaded_file, file_type: str) -> List[Document]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            if file_type.lower() == 'pdf':
                documents = self.load_pdf(tmp_file_path)
            elif file_type.lower() in ['docx', 'doc']:
                documents = self.load_docx(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            for doc in documents:
                doc.metadata['filename'] = uploaded_file.name
            
            return documents
        finally:
            os.unlink(tmp_file_path)
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        return self.text_splitter.split_documents(documents)
    
    def process_documents(self, uploaded_files) -> List[Document]:
        all_documents = []
        
        for uploaded_file in uploaded_files:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension in ['pdf', 'docx', 'doc']:
                documents = self.process_uploaded_file(uploaded_file, file_extension)
                all_documents.extend(documents)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        return self.split_documents(all_documents)